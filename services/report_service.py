import io
from datetime import date, datetime
import calendar
from sqlalchemy import func
from extensions import db
from models.expense import Expense
from models.income import Income
from models.budget import Budget
from models.goal import Goal
from models.account import Account
from models.investment import Investment
from models.alert import FinancialAlert
from services.spending_analysis import calculate_financial_health_score
from services.alert_service import get_user_alerts

# Export Libraries
import reportlab
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, KeepTogether

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter


def parse_month_year(month_input, year_input):
    """
    Parses and sanitizes month and year input into valid integer month (1-12) and year.
    Defaults to current month and year if invalid or missing.
    """
    today = date.today()
    try:
        if isinstance(month_input, str):
            if month_input.isdigit():
                month = int(month_input)
            else:
                # Try month names e.g. "January" or "Jan"
                try:
                    month = datetime.strptime(month_input[:3], "%b").month
                except ValueError:
                    month = today.month
        elif isinstance(month_input, int):
            month = month_input
        else:
            month = today.month

        if not (1 <= month <= 12):
            month = today.month
    except Exception:
        month = today.month

    try:
        year = int(year_input) if year_input else today.year
        if year < 2000 or year > 2100:
            year = today.year
    except Exception:
        year = today.year

    return month, year


def get_monthly_financial_report(user_id, month_input=None, year_input=None):
    """
    Generates a dynamic Monthly Financial Report for the authenticated user for the specified month & year.
    Uses actual database records restricted strictly to user_id.
    """
    month, year = parse_month_year(month_input, year_input)
    month_name = calendar.month_name[month]
    num_days = calendar.monthrange(year, month)[1]
    start_date = date(year, month, 1)
    end_date = date(year, month, num_days)

    # 1. Income for selected month
    monthly_incomes = Income.query.filter(
        Income.user_id == user_id,
        Income.income_date >= start_date,
        Income.income_date <= end_date
    ).order_by(Income.income_date.asc()).all()

    total_income = db.session.query(func.sum(Income.amount)).filter(
        Income.user_id == user_id,
        Income.income_date >= start_date,
        Income.income_date <= end_date
    ).scalar() or 0.0

    # 2. Expenses for selected month
    monthly_expenses = Expense.query.filter(
        Expense.user_id == user_id,
        Expense.expense_date >= start_date,
        Expense.expense_date <= end_date
    ).order_by(Expense.expense_date.asc()).all()

    total_expenses = db.session.query(func.sum(Expense.amount)).filter(
        Expense.user_id == user_id,
        Expense.expense_date >= start_date,
        Expense.expense_date <= end_date
    ).scalar() or 0.0

    net_savings = total_income - total_expenses
    savings_rate = round((net_savings / total_income) * 100, 1) if total_income > 0 else 0.0

    # 3. Category Breakdown for selected month (SQL Grouped Aggregation)
    category_data = (
        db.session.query(Expense.category, func.sum(Expense.amount))
        .filter(
            Expense.user_id == user_id,
            Expense.expense_date >= start_date,
            Expense.expense_date <= end_date
        )
        .group_by(Expense.category)
        .order_by(func.sum(Expense.amount).desc())
        .all()
    )

    category_breakdown = [
        {
            "category": cat,
            "amount": float(amt),
            "percentage": round((float(amt) / total_expenses) * 100, 1) if total_expenses > 0 else 0.0
        }
        for cat, amt in category_data
    ]

    # 4. Budget for selected month
    # First match exact month & year if exists, else match active budget
    month_str = calendar.month_name[month]
    budget_obj = Budget.query.filter(
        Budget.user_id == user_id,
        Budget.year == year
    ).filter(
        (Budget.month == month_str) | (Budget.month == f"{month:02d}") | (Budget.month.startswith(month_str[:3]))
    ).first()

    if not budget_obj:
        budget_obj = Budget.query.filter_by(user_id=user_id).order_by(Budget.created_at.desc()).first()

    budget_amount = budget_obj.monthly_budget if budget_obj else 0.0
    budget_remaining = budget_amount - total_expenses if budget_amount > 0 else 0.0
    budget_utilization = round((total_expenses / budget_amount) * 100, 1) if budget_amount > 0 else 0.0

    # 5. Accounts & Investments Overview
    accounts = Account.query.filter_by(user_id=user_id).all()
    total_account_balance = sum(acc.balance for acc in accounts)

    investments = Investment.query.filter_by(user_id=user_id).all()
    total_invested = sum(inv.invested_amount for inv in investments)
    total_current_investment_val = sum(inv.current_value for inv in investments)
    net_worth = total_account_balance + total_current_investment_val

    # 6. Goals Activity & Linked Expenses for selected month
    goals = Goal.query.filter_by(user_id=user_id).all()
    goal_activity = []
    for g in goals:
        linked_month_exp = [
            e for e in monthly_expenses if e.goal_id == g.id
        ]
        month_exp_sum = sum(e.amount for e in linked_month_exp)
        prog_pct = round((g.current_amount / g.target_amount) * 100, 1) if g.target_amount > 0 else 0.0
        goal_activity.append({
            "goal": g,
            "goal_name": g.goal_name,
            "target_amount": g.target_amount,
            "current_amount": g.current_amount,
            "progress_percentage": prog_pct,
            "monthly_expenses_count": len(linked_month_exp),
            "monthly_expenses_sum": month_exp_sum,
            "status": g.status
        })

    # 7. Financial Health Score & Alerts
    health_score = calculate_financial_health_score(user_id)
    alerts = get_user_alerts(user_id, include_read=True)

    return {
        "user_id": user_id,
        "month": month,
        "month_name": month_name,
        "year": year,
        "period_label": f"{month_name} {year}",
        "start_date": start_date,
        "end_date": end_date,
        "total_income": total_income,
        "total_expenses": total_expenses,
        "net_savings": net_savings,
        "savings_rate": savings_rate,
        "monthly_incomes": monthly_incomes,
        "monthly_expenses": monthly_expenses,
        "category_breakdown": category_breakdown,
        "budget": budget_obj,
        "budget_amount": budget_amount,
        "budget_remaining": budget_remaining,
        "budget_utilization": budget_utilization,
        "accounts": accounts,
        "total_account_balance": total_account_balance,
        "investments": investments,
        "total_invested": total_invested,
        "total_current_investment_val": total_current_investment_val,
        "net_worth": net_worth,
        "goals": goals,
        "goal_activity": goal_activity,
        "health_score": health_score,
        "alerts": alerts
    }


def get_goal_progress_report(user_id):
    """
    Generates detailed Goal Progress Report for the authenticated user.
    """
    goals = Goal.query.filter_by(user_id=user_id).order_by(Goal.created_at.desc()).all()
    all_expenses = Expense.query.filter_by(user_id=user_id).all()
    all_budgets = Budget.query.filter_by(user_id=user_id).all()

    report_goals = []
    total_target = sum(g.target_amount for g in goals)
    total_saved = sum(g.current_amount for g in goals)
    overall_progress = round((total_saved / total_target) * 100, 1) if total_target > 0 else 0.0

    for g in goals:
        remaining = max(0.0, g.target_amount - g.current_amount)
        prog_pct = round((g.current_amount / g.target_amount) * 100, 1) if g.target_amount > 0 else 0.0
        linked_expenses = [e for e in all_expenses if e.goal_id == g.id]
        linked_exp_total = sum(e.amount for e in linked_expenses)
        linked_budget = next((b for b in all_budgets if b.goal_id == g.id), None)

        report_goals.append({
            "id": g.id,
            "goal_name": g.goal_name,
            "goal_type": g.goal_type,
            "category": g.category,
            "priority": g.priority,
            "target_amount": g.target_amount,
            "current_amount": g.current_amount,
            "remaining_amount": remaining,
            "progress_percentage": prog_pct,
            "target_date": g.target_date,
            "target_date_str": g.target_date.strftime("%d %b %Y") if g.target_date else "N/A",
            "status": g.status,
            "linked_budget": linked_budget,
            "linked_expenses": linked_expenses,
            "linked_expenses_count": len(linked_expenses),
            "linked_expenses_total": linked_exp_total
        })

    return {
        "goals": report_goals,
        "total_goals_count": len(goals),
        "total_target": total_target,
        "total_saved": total_saved,
        "overall_progress": overall_progress
    }


# ==============================================================================
# PDF EXPORT GENERATOR (ReportLab)
# ==============================================================================

def generate_pdf_report(user_name, monthly_report, goal_report):
    """
    Builds a publication-quality PDF financial report in memory.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=36,
        leftMargin=36,
        topMargin=36,
        bottomMargin=36
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontSize=22,
        leading=26,
        textColor=colors.HexColor("#1E3A8A"),
        alignment=0,
        spaceAfter=4
    )

    subtitle_style = ParagraphStyle(
        'DocSubTitle',
        parent=styles['Normal'],
        fontSize=11,
        leading=14,
        textColor=colors.HexColor("#4B5563"),
        spaceAfter=15
    )

    h2_style = ParagraphStyle(
        'DocH2',
        parent=styles['Heading2'],
        fontSize=14,
        leading=18,
        textColor=colors.HexColor("#1F2937"),
        spaceBefore=12,
        spaceAfter=8
    )

    body_style = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#374151")
    )

    table_header_style = ParagraphStyle(
        'TableHeader',
        parent=styles['Normal'],
        fontSize=9,
        leading=11,
        fontName='Helvetica-Bold',
        textColor=colors.white
    )

    story = []

    # Header section
    story.append(Paragraph("FinSight — Personal Financial Report", title_style))
    story.append(Paragraph(
        f"<b>Report Period:</b> {monthly_report['period_label']} &nbsp;|&nbsp; "
        f"<b>User:</b> {user_name} &nbsp;|&nbsp; "
        f"<b>Generated:</b> {date.today().strftime('%B %d, %Y')}",
        subtitle_style
    ))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#2563EB"), spaceAfter=12))

    # Executive Summary Cards Table
    summary_data = [
        [
            Paragraph("<b>Total Income</b>", body_style),
            Paragraph("<b>Total Expenses</b>", body_style),
            Paragraph("<b>Net Savings</b>", body_style),
            Paragraph("<b>Health Score</b>", body_style)
        ],
        [
            Paragraph(f"<font size=12 color='#16A34A'><b>₹{monthly_report['total_income']:,.2f}</b></font>", body_style),
            Paragraph(f"<font size=12 color='#DC2626'><b>₹{monthly_report['total_expenses']:,.2f}</b></font>", body_style),
            Paragraph(f"<font size=12 color='#2563EB'><b>₹{monthly_report['net_savings']:,.2f}</b></font>", body_style),
            Paragraph(f"<font size=12 color='#7C3AED'><b>{monthly_report['health_score']['score']}/100 ({monthly_report['health_score']['status_label']})</b></font>", body_style)
        ]
    ]
    summary_table = Table(summary_data, colWidths=[130, 130, 130, 150])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#F3F4F6")),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#D1D5DB")),
        ('TOPPADDING', (0,0), (-1,-1), 8),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(summary_table)
    story.append(Spacer(1, 14))

    # Category Expense Breakdown Table
    story.append(Paragraph("Category-wise Expense Breakdown", h2_style))
    cat_rows = [
        [Paragraph("Category", table_header_style), Paragraph("Amount (₹)", table_header_style), Paragraph("Share (%)", table_header_style)]
    ]
    for cat in monthly_report['category_breakdown']:
        cat_rows.append([
            Paragraph(cat['category'], body_style),
            Paragraph(f"₹{cat['amount']:,.2f}", body_style),
            Paragraph(f"{cat['percentage']}%", body_style)
        ])
    if len(cat_rows) == 1:
        cat_rows.append([Paragraph("No expenses recorded for this month.", body_style), Paragraph("-", body_style), Paragraph("-", body_style)])

    cat_table = Table(cat_rows, colWidths=[240, 150, 150])
    cat_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#1E3A8A")),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#E5E7EB")),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor("#F9FAFB")]),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
    ]))
    story.append(cat_table)
    story.append(Spacer(1, 14))

    # Budget & Net Worth Overview
    story.append(Paragraph("Budget & Net Worth Overview", h2_style))
    budget_val = f"₹{monthly_report['budget_amount']:,.2f}" if monthly_report['budget_amount'] > 0 else "No Budget Set"
    util_val = f"{monthly_report['budget_utilization']}%" if monthly_report['budget_amount'] > 0 else "N/A"
    bw_data = [
        [Paragraph("Monthly Budget", body_style), Paragraph(budget_val, body_style), Paragraph("Budget Utilization", body_style), Paragraph(util_val, body_style)],
        [Paragraph("Account Balances", body_style), Paragraph(f"₹{monthly_report['total_account_balance']:,.2f}", body_style), Paragraph("Investments Value", body_style), Paragraph(f"₹{monthly_report['total_current_investment_val']:,.2f}", body_style)],
        [Paragraph("Estimated Net Worth", body_style), Paragraph(f"₹{monthly_report['net_worth']:,.2f}", body_style), Paragraph("Net Savings Margin", body_style), Paragraph(f"{monthly_report['savings_rate']}%", body_style)]
    ]
    bw_table = Table(bw_data, colWidths=[130, 140, 130, 140])
    bw_table.setStyle(TableStyle([
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#D1D5DB")),
        ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#F3F4F6")),
        ('BACKGROUND', (2,0), (2,-1), colors.HexColor("#F3F4F6")),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
    ]))
    story.append(bw_table)
    story.append(Spacer(1, 14))

    # Goal Progress Report Table
    story.append(Paragraph("Goal Progress Summary", h2_style))
    goal_rows = [
        [Paragraph("Goal Name", table_header_style), Paragraph("Target (₹)", table_header_style), Paragraph("Saved (₹)", table_header_style), Paragraph("Progress", table_header_style), Paragraph("Target Date", table_header_style), Paragraph("Status", table_header_style)]
    ]
    for g in goal_report['goals']:
        goal_rows.append([
            Paragraph(g['goal_name'], body_style),
            Paragraph(f"₹{g['target_amount']:,.2f}", body_style),
            Paragraph(f"₹{g['current_amount']:,.2f}", body_style),
            Paragraph(f"{g['progress_percentage']}%", body_style),
            Paragraph(g['target_date_str'], body_style),
            Paragraph(g['status'], body_style)
        ])
    if len(goal_rows) == 1:
        goal_rows.append([Paragraph("No financial goals set.", body_style), Paragraph("-", body_style), Paragraph("-", body_style), Paragraph("-", body_style), Paragraph("-", body_style), Paragraph("-", body_style)])

    goal_table = Table(goal_rows, colWidths=[130, 85, 85, 70, 90, 80])
    goal_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#1D4ED8")),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#E5E7EB")),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor("#F9FAFB")]),
        ('TOPPADDING', (0,0), (-1,-1), 5),
        ('BOTTOMPADDING', (0,0), (-1,-1), 5),
    ]))
    story.append(goal_table)

    doc.build(story)
    buffer.seek(0)
    return buffer


# ==============================================================================
# WORD / DOCX EXPORT GENERATOR (python-docx)
# ==============================================================================

def generate_docx_report(user_name, monthly_report, goal_report):
    """
    Builds a cleanly styled Word (.docx) financial report in memory.
    """
    doc = docx.Document()

    # Title
    title = doc.add_heading("FinSight — Personal Financial Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Meta paragraph
    meta = doc.add_paragraph()
    meta.add_run(f"Report Period: {monthly_report['period_label']}  |  User: {user_name}  |  Date: {date.today().strftime('%B %d, %Y')}\n").italic = True

    # Executive Summary Heading
    doc.add_heading("1. Executive Financial Summary", level=1)
    
    summary_table = doc.add_table(rows=2, cols=4)
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = summary_table.rows[0].cells
    hdr_cells[0].text = "Total Income"
    hdr_cells[1].text = "Total Expenses"
    hdr_cells[2].text = "Net Savings"
    hdr_cells[3].text = "Health Score"

    val_cells = summary_table.rows[1].cells
    val_cells[0].text = f"₹{monthly_report['total_income']:,.2f}"
    val_cells[1].text = f"₹{monthly_report['total_expenses']:,.2f}"
    val_cells[2].text = f"₹{monthly_report['net_savings']:,.2f}"
    val_cells[3].text = f"{monthly_report['health_score']['score']}/100"

    doc.add_paragraph()

    # Category Breakdown
    doc.add_heading("2. Expense Category Breakdown", level=1)
    cat_table = doc.add_table(rows=1, cols=3)
    cat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_hdr = cat_table.rows[0].cells
    c_hdr[0].text = "Category"
    c_hdr[1].text = "Amount (₹)"
    c_hdr[2].text = "Percentage Share"

    for cat in monthly_report['category_breakdown']:
        row_cells = cat_table.add_row().cells
        row_cells[0].text = cat['category']
        row_cells[1].text = f"₹{cat['amount']:,.2f}"
        row_cells[2].text = f"{cat['percentage']}%"

    doc.add_paragraph()

    # Budget & Goals
    doc.add_heading("3. Budget & Goal Progress", level=1)
    b_p = doc.add_paragraph()
    b_p.add_run(f"Monthly Budget Limit: ₹{monthly_report['budget_amount']:,.2f}\n")
    b_p.add_run(f"Budget Utilization: {monthly_report['budget_utilization']}%\n")
    b_p.add_run(f"Net Savings Rate: {monthly_report['savings_rate']}%\n")

    goal_table = doc.add_table(rows=1, cols=5)
    g_hdr = goal_table.rows[0].cells
    g_hdr[0].text = "Goal Name"
    g_hdr[1].text = "Target Amount"
    g_hdr[2].text = "Current Savings"
    g_hdr[3].text = "Progress"
    g_hdr[4].text = "Status"

    for g in goal_report['goals']:
        row_cells = goal_table.add_row().cells
        row_cells[0].text = g['goal_name']
        row_cells[1].text = f"₹{g['target_amount']:,.2f}"
        row_cells[2].text = f"₹{g['current_amount']:,.2f}"
        row_cells[3].text = f"{g['progress_percentage']}%"
        row_cells[4].text = g['status']

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ==============================================================================
# EXCEL EXPORT GENERATOR (openpyxl - 8 Structured Sheets)
# ==============================================================================

def generate_excel_report(user_name, monthly_report, goal_report):
    """
    Builds a multi-sheet structured Excel workbook containing 8 sheets.
    """
    wb = openpyxl.Workbook()
    
    # Styles
    header_fill = PatternFill(start_color="1E3A8A", end_color="1E3A8A", fill_type="solid")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    bold_font = Font(name="Calibri", size=11, bold=True)
    title_font = Font(name="Calibri", size=16, bold=True, color="1E3A8A")
    thin_border = Border(
        left=Side(style='thin', color='D1D5DB'),
        right=Side(style='thin', color='D1D5DB'),
        top=Side(style='thin', color='D1D5DB'),
        bottom=Side(style='thin', color='D1D5DB')
    )

    def style_header(ws, row=1):
        for cell in ws[row]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")

    def auto_fit_columns(ws):
        for col in ws.columns:
            max_len = 0
            col_letter = get_column_letter(col[0].column)
            for cell in col:
                val_str = str(cell.value or '')
                max_len = max(max_len, len(val_str))
            ws.column_dimensions[col_letter].width = max(max_len + 3, 12)

    # -------------------------------------------------------------
    # Sheet 1: Summary
    # -------------------------------------------------------------
    ws_summary = wb.active
    ws_summary.title = "Summary"

    ws_summary.append(["FinSight — Financial Summary Report"])
    ws_summary.cell(row=1, column=1).font = title_font
    ws_summary.append([f"Report Period: {monthly_report['period_label']}", f"User: {user_name}", f"Date: {date.today().strftime('%Y-%m-%d')}"])
    ws_summary.append([])

    ws_summary.append(["Metric", "Amount (₹) / Value"])
    style_header(ws_summary, row=4)

    metrics = [
        ("Total Income", monthly_report['total_income']),
        ("Total Expenses", monthly_report['total_expenses']),
        ("Net Savings", monthly_report['net_savings']),
        ("Savings Rate (%)", f"{monthly_report['savings_rate']}%"),
        ("Monthly Budget Amount", monthly_report['budget_amount']),
        ("Budget Utilization (%)", f"{monthly_report['budget_utilization']}%"),
        ("Total Account Balances", monthly_report['total_account_balance']),
        ("Investments Value", monthly_report['total_current_investment_val']),
        ("Estimated Net Worth", monthly_report['net_worth']),
        ("Financial Health Score", f"{monthly_report['health_score']['score']}/100 ({monthly_report['health_score']['status_label']})")
    ]

    for m, val in metrics:
        ws_summary.append([m, val])

    auto_fit_columns(ws_summary)

    # -------------------------------------------------------------
    # Sheet 2: Income
    # -------------------------------------------------------------
    ws_inc = wb.create_sheet(title="Income")
    ws_inc.append(["ID", "Title", "Source", "Amount (₹)", "Income Date", "Description"])
    style_header(ws_inc, row=1)

    for inc in monthly_report['monthly_incomes']:
        ws_inc.append([
            inc.id, inc.title, inc.source, inc.amount,
            inc.income_date.strftime("%Y-%m-%d") if inc.income_date else "",
            inc.description or ""
        ])
    auto_fit_columns(ws_inc)

    # -------------------------------------------------------------
    # Sheet 3: Expenses
    # -------------------------------------------------------------
    ws_exp = wb.create_sheet(title="Expenses")
    ws_exp.append(["ID", "Title", "Category", "Amount (₹)", "Payment Method", "Expense Date", "Goal ID", "Description"])
    style_header(ws_exp, row=1)

    for exp in monthly_report['monthly_expenses']:
        ws_exp.append([
            exp.id, exp.title, exp.category, exp.amount, exp.payment_method,
            exp.expense_date.strftime("%Y-%m-%d") if exp.expense_date else "",
            exp.goal_id or "", exp.description or ""
        ])
    auto_fit_columns(ws_exp)

    # -------------------------------------------------------------
    # Sheet 4: Budget
    # -------------------------------------------------------------
    ws_bgt = wb.create_sheet(title="Budget")
    ws_bgt.append(["ID", "Month", "Year", "Monthly Budget (₹)", "Total Spent (₹)", "Utilization (%)"])
    style_header(ws_bgt, row=1)

    if monthly_report['budget']:
        b = monthly_report['budget']
        ws_bgt.append([
            b.id, b.month, b.year, b.monthly_budget,
            monthly_report['total_expenses'], monthly_report['budget_utilization']
        ])
    auto_fit_columns(ws_bgt)

    # -------------------------------------------------------------
    # Sheet 5: Goals
    # -------------------------------------------------------------
    ws_goals = wb.create_sheet(title="Goals")
    ws_goals.append(["ID", "Goal Name", "Category", "Target Amount (₹)", "Current Saved (₹)", "Progress (%)", "Target Date", "Status"])
    style_header(ws_goals, row=1)

    for g in goal_report['goals']:
        ws_goals.append([
            g['id'], g['goal_name'], g['category'], g['target_amount'],
            g['current_amount'], g['progress_percentage'], g['target_date_str'], g['status']
        ])
    auto_fit_columns(ws_goals)

    # -------------------------------------------------------------
    # Sheet 6: Goal Expenses
    # -------------------------------------------------------------
    ws_gexp = wb.create_sheet(title="Goal Expenses")
    ws_gexp.append(["Goal ID", "Goal Name", "Linked Expenses Count", "Total Linked Expenses (₹)"])
    style_header(ws_gexp, row=1)

    for g in goal_report['goals']:
        ws_gexp.append([
            g['id'], g['goal_name'], g['linked_expenses_count'], g['linked_expenses_total']
        ])
    auto_fit_columns(ws_gexp)

    # -------------------------------------------------------------
    # Sheet 7: Investments
    # -------------------------------------------------------------
    ws_inv = wb.create_sheet(title="Investments")
    ws_inv.append(["ID", "Instrument Name", "Asset Type", "Quantity", "Invested Amount (₹)", "Current Value (₹)", "Purchase Date"])
    style_header(ws_inv, row=1)

    for inv in monthly_report['investments']:
        ws_inv.append([
            inv.id, inv.instrument_name, inv.asset_type, inv.quantity,
            inv.invested_amount, inv.current_value,
            inv.purchase_date.strftime("%Y-%m-%d") if inv.purchase_date else ""
        ])
    auto_fit_columns(ws_inv)

    # -------------------------------------------------------------
    # Sheet 8: Alerts
    # -------------------------------------------------------------
    ws_alr = wb.create_sheet(title="Alerts")
    ws_alr.append(["ID", "Type", "Title", "Severity", "Message", "Is Read", "Date"])
    style_header(ws_alr, row=1)

    for al in monthly_report['alerts']:
        ws_alr.append([
            al.id, al.alert_type, al.title, al.severity, al.message,
            "Yes" if al.is_read else "No",
            al.created_at.strftime("%Y-%m-%d %H:%M") if al.created_at else ""
        ])
    auto_fit_columns(ws_alr)

    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer
